"""
ArgusHR RAG Engine with PostgreSQL & Pgvector
Enhanced RAG implementation with cloud-compatible persistent vector storage.
"""
import os
import re
import sys
import time
import json
import hashlib
import logging
import threading

try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

from pathlib import Path
from typing import List, Dict, Tuple, Optional, Generator
from datetime import datetime

import numpy as np
import tiktoken
import voyageai
from dotenv import load_dotenv
import pandas as pd
import pdfplumber

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv(override=True)

# Try loading psycopg2
try:
    import psycopg2
    from psycopg2.extras import execute_values
except ImportError:
    logger.error("❌ psycopg2-binary not installed. Please run: pip install psycopg2-binary")
    raise ImportError("psycopg2-binary required for RAGEngine PostgreSQL connection")

# ---------------- CONFIGURATION ----------------
class Config:
    # Embedding settings
    EMBED_MODEL = "voyage-4-large"
    RERANK_MODEL = "rerank-2.5"
    
    # Chunking settings
    MAX_TOKENS = 800
    OVERLAP_TOKENS = 120
    
    # Retrieval settings
    TOP_K = 25
    RERANK_TOP_K = 3
    
    # API limits
    MAX_BATCH_SIZE = 100
    MAX_BATCH_TOKENS = 90000
    WAIT_TIME = 0.3

    # Indexing memory guard: embed + insert this many chunks at a time so peak
    # process memory stays flat regardless of how many/large the files are.
    INDEX_SLICE_SIZE = 100
    
    # LLM settings
    MAX_CONTEXT_TOKENS = 12000
    LLM_MODEL = "llama-3.3-70b-versatile"
    
    # Directories
    DOCUMENTS_DIR = "documents"
    
    @classmethod
    def set_documents_dir(cls, path: str):
        """Set documents directory to a specific path or subfolder."""
        cls.DOCUMENTS_DIR = path

config = Config()

# ---------------- TOKENIZER ----------------
enc = tiktoken.get_encoding("cl100k_base")

def count_tokens(text: str) -> int:
    """Count tokens in text."""
    return len(enc.encode(text))

# ---------------- CUSTOM EXCEPTIONS ----------------
class RAGError(Exception):
    """Base RAG exception."""
    pass

class DocumentProcessingError(RAGError):
    """Error processing a document."""
    pass

class EmbeddingError(RAGError):
    """Error generating embeddings."""
    pass

class QueryError(RAGError):
    """Error during query."""
    pass

# ---------------- FILE UTILITIES ----------------
def get_file_hash(file_path: str) -> str:
    """Generate MD5 hash for change detection."""
    hasher = hashlib.md5()
    try:
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hasher.update(chunk)
        return hasher.hexdigest()
    except IOError as e:
        logger.error(f"Could not hash file {file_path}: {e}")
        raise DocumentProcessingError(f"Cannot read file: {file_path}")

def discover_files(directory: str) -> List[Dict]:
    """Discover all supported files in directory."""
    supported_extensions = {'.xlsx', '.xls', '.pdf', '.md', '.txt', '.docx', '.csv', '.tsv'}
    files = []
    
    doc_path = Path(directory)
    if not doc_path.exists():
        logger.info(f"Creating {directory}/ folder...")
        doc_path.mkdir(parents=True, exist_ok=True)
        return []
    
    for file_path in doc_path.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            try:
                relative_path = file_path.relative_to(doc_path)
            except ValueError:
                relative_path = file_path.name

            files.append({
                'path': str(file_path),
                'relative_path': str(relative_path),
                'name': file_path.name,
                'type': file_path.suffix.lower(),
                'hash': get_file_hash(str(file_path)),
                'size': file_path.stat().st_size,
                'modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
            })
    
    return sorted(files, key=lambda x: x['name'])

# ---------------- DOCUMENT PROCESSORS ----------------
def _render_sheet_rows(df) -> List[str]:
    """One text line per row, dropping empty cells (read with header=None upstream)."""
    rows: List[str] = []
    for _, r in df.iterrows():
        cells = [str(v).strip() for v in r.tolist()]
        cells = [c for c in cells if c and c.lower() != "nan"]
        if cells:
            rows.append(" | ".join(cells))
    return rows


def _sheet_to_chunks(sheet_name: str, df, overlap_rows: int = 2) -> List[str]:
    """Token-bounded, self-describing chunks with row overlap; oversized rows hard-split."""
    rows = _render_sheet_rows(df)
    if not rows:
        return []
    prefix = f"## Sheet: {sheet_name}\n"
    budget = max(200, config.MAX_TOKENS - count_tokens(prefix))
    out, cur, cur_tokens = [], [], 0

    def flush(keep_overlap):
        nonlocal cur, cur_tokens
        if cur:
            out.append(prefix + "\n".join(cur))
            cur = cur[-overlap_rows:] if (keep_overlap and overlap_rows) else []
            cur_tokens = sum(count_tokens(x) for x in cur)

    for row in rows:
        t = count_tokens(row)
        if t > budget:
            flush(False)
            toks = enc.encode(row)
            for s in range(0, len(toks), budget):
                out.append(prefix + enc.decode(toks[s:s + budget]))
            continue
        if cur_tokens + t > budget and cur:
            flush(True)
        cur.append(row)
        cur_tokens += t
    flush(False)
    return out


def process_excel(file_path: str) -> List[Dict]:
    """
    Process an Excel workbook: every sheet, every row (header=None so nothing is
    lost/garbled), chunked with overlap. Each chunk carries sheet + chunk_index.
    """
    chunks: List[Dict] = []
    try:
        with pd.ExcelFile(file_path) as excel_file:
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
                df = df.dropna(how="all")
                if df.empty:
                    continue
                for ci, part in enumerate(_sheet_to_chunks(sheet_name, df)):
                    chunks.append({
                        "text": part,
                        "metadata": {
                            "source": file_path,
                            "source_name": Path(file_path).name,
                            "sheet": sheet_name,
                            "type": "excel",
                            "chunk_index": ci,
                            "indexed_at": datetime.now().isoformat(),
                        },
                    })
            logger.info(
                f"✓ Processed {len(excel_file.sheet_names)} sheet(s) from "
                f"{Path(file_path).name}, created {len(chunks)} chunks"
            )
    except Exception as e:
        logger.error(f"Error processing Excel {file_path}: {e}")
        raise DocumentProcessingError(f"Excel processing failed: {e}")
    return chunks

def process_pdf(file_path: str) -> List[Dict]:
    """Process PDF file with enhanced metadata."""
    chunks = []
    
    try:
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                
                if not text or len(text.strip()) < 50:
                    continue
                
                page_text = f"## Page {page_num} of {total_pages}\n\n{text}"
                
                if count_tokens(page_text) <= config.MAX_TOKENS:
                    chunks.append({
                        'text': page_text,
                        'metadata': {
                            'source': file_path,
                            'source_name': Path(file_path).name,
                            'page': page_num,
                            'total_pages': total_pages,
                            'type': 'pdf',
                            'indexed_at': datetime.now().isoformat()
                        }
                    })
                else:
                    paragraphs = text.split('\n\n')
                    current_chunk = f"## Page {page_num} of {total_pages}\n\n"
                    
                    for para in paragraphs:
                        if count_tokens(current_chunk + para) > config.MAX_TOKENS:
                            if current_chunk.strip():
                                chunks.append({
                                    'text': current_chunk.strip(),
                                    'metadata': {
                                        'source': file_path,
                                        'source_name': Path(file_path).name,
                                        'page': page_num,
                                        'total_pages': total_pages,
                                        'type': 'pdf',
                                        'indexed_at': datetime.now().isoformat()
                                    }
                                })
                            current_chunk = f"## Page {page_num} of {total_pages}\n\n{para}\n\n"
                        else:
                            current_chunk += para + "\n\n"
                    
                    if current_chunk.strip():
                        chunks.append({
                            'text': current_chunk.strip(),
                            'metadata': {
                                'source': file_path,
                                'source_name': Path(file_path).name,
                                'page': page_num,
                                'total_pages': total_pages,
                                'type': 'pdf',
                                'indexed_at': datetime.now().isoformat()
                            }
                        })
            
            logger.info(f"✓ Processed {total_pages} pages from {Path(file_path).name}, created {len(chunks)} chunks")
            
    except Exception as e:
        logger.error(f"Error processing PDF {file_path}: {e}")
        raise DocumentProcessingError(f"PDF processing failed: {e}")
    
    return chunks

def split_markdown_blocks(md: str) -> List[str]:
    """Split markdown into semantic blocks."""
    blocks = []
    buf = []
    in_code = False

    lines = md.splitlines()
    for line in lines:
        if line.strip().startswith("```"):
            in_code = not in_code
            buf.append(line)
            continue

        if in_code:
            buf.append(line)
            continue

        if re.match(r"^#{1,6}\s+", line):
            if buf:
                blocks.append("\n".join(buf).strip())
                buf = []
            buf.append(line)
            continue

        if "|" in line and re.match(r"^\s*\|.*\|\s*$", line):
            buf.append(line)
            continue

        if re.match(r"^\s*[-*+]\s+", line):
            buf.append(line)
            continue

        if line.strip() == "":
            if buf:
                blocks.append("\n".join(buf).strip())
                buf = []
            continue

        buf.append(line)

    if buf:
        blocks.append("\n".join(buf).strip())

    return [b for b in blocks if b.strip()]

def chunk_markdown_blocks(blocks: List[str], max_tokens: int, overlap_tokens: int) -> List[str]:
    """Chunk markdown blocks with overlap."""
    chunks = []
    current = []
    current_tokens = 0

    def flush_with_overlap():
        nonlocal current, current_tokens
        if not current:
            return
        chunk_text = "\n\n".join(current).strip()
        chunks.append(chunk_text)

        if overlap_tokens > 0:
            tokens = enc.encode(chunk_text)
            overlap = tokens[-overlap_tokens:] if len(tokens) > overlap_tokens else tokens
            overlap_text = enc.decode(overlap)
            current = [overlap_text]
            current_tokens = count_tokens(overlap_text)
        else:
            current = []
            current_tokens = 0

    for block in blocks:
        block_tokens = count_tokens(block)

        if block_tokens > max_tokens:
            if current:
                flush_with_overlap()

            tokens = enc.encode(block)
            start = 0
            while start < len(tokens):
                end = min(start + max_tokens, len(tokens))
                chunk = enc.decode(tokens[start:end])
                chunks.append(chunk)
                start = end - overlap_tokens if overlap_tokens > 0 and end - overlap_tokens > start else end
            current = []
            current_tokens = 0
            continue

        if current_tokens + block_tokens <= max_tokens:
            current.append(block)
            current_tokens += block_tokens
        else:
            flush_with_overlap()
            current.append(block)
            current_tokens += block_tokens

    if current:
        flush_with_overlap()

    return chunks

def process_markdown(file_path: str) -> List[Dict]:
    """Process Markdown file with header extraction."""
    chunks = []
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract document title
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        doc_title = title_match.group(1) if title_match else Path(file_path).stem
        
        blocks = split_markdown_blocks(content)
        chunk_texts = chunk_markdown_blocks(blocks, config.MAX_TOKENS, config.OVERLAP_TOKENS)
        
        for i, text in enumerate(chunk_texts):
            # Extract section header if present
            header_match = re.search(r'^#{1,6}\s+(.+)$', text, re.MULTILINE)
            section = header_match.group(1) if header_match else None
            
            chunks.append({
                'text': text,
                'metadata': {
                    'source': file_path,
                    'source_name': Path(file_path).name,
                    'chunk_id': i,
                    'type': 'markdown',
                    'title': doc_title,
                    'section': section or "",
                    'indexed_at': datetime.now().isoformat()
                }
            })
        
        logger.info(f"✓ Processed {Path(file_path).name}, created {len(chunks)} chunks")
        
    except Exception as e:
        logger.error(f"Error processing Markdown {file_path}: {e}")
        raise DocumentProcessingError(f"Markdown processing failed: {e}")
    
    return chunks

def process_text(file_path: str) -> List[Dict]:
    """Process plain text file."""
    chunks = []
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        paragraphs = content.split('\n\n')
        current_chunk = ""
        chunk_id = 0
        
        for para in paragraphs:
            if count_tokens(current_chunk + para) > config.MAX_TOKENS:
                if current_chunk.strip():
                    chunks.append({
                        'text': current_chunk.strip(),
                        'metadata': {
                            'source': file_path,
                            'source_name': Path(file_path).name,
                            'chunk_id': chunk_id,
                            'type': 'text',
                            'indexed_at': datetime.now().isoformat()
                        }
                    })
                    chunk_id += 1
                current_chunk = para + "\n\n"
            else:
                current_chunk += para + "\n\n"
        
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'metadata': {
                    'source': file_path,
                    'source_name': Path(file_path).name,
                    'chunk_id': chunk_id,
                    'type': 'text',
                    'indexed_at': datetime.now().isoformat()
                }
            })
        
        logger.info(f"✓ Processed {Path(file_path).name}, created {len(chunks)} chunks")
        
    except Exception as e:
        logger.error(f"Error processing text {file_path}: {e}")
        raise DocumentProcessingError(f"Text processing failed: {e}")
    
    return chunks

def process_docx(file_path: str) -> List[Dict]:
    """Process Microsoft Word .docx file."""
    chunks = []
    try:
        import docx
        with open(file_path, "rb") as f:
            doc = docx.Document(f)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract tables if present
        for table in doc.tables:
            if not table.rows:
                continue
            # Get headers from the first row of the table
            headers = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
            headers_str = " | ".join([h for h in headers if h])
            
            # Process each row and append the headers context
            start_row = 1 if len(table.rows) > 1 else 0
            for row in table.rows[start_row:]:
                row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                # Filter out empty rows
                if any(row_text):
                    # Format as: Table [Header1 | Header2]: Val1 | Val2
                    full_text.append(f"Table [{headers_str}]: " + " | ".join(row_text))
                    
        content = "\n\n".join(full_text)
        
        paragraphs = content.split('\n\n')
        current_chunk = ""
        chunk_id = 0
        current_section = ""
        
        for para in paragraphs:
            para_strip = para.strip()
            # Detect section heading pattern
            if re.match(r'^(?:SECTION\s+\d+|Section\s+\d+|\d+\.\d+|\d+\.\d+\.\d+)\b', para_strip) or \
               (para_strip.isupper() and len(para_strip) < 60) or \
               para_strip.startswith("###") or para_strip.startswith("##"):
                current_section = para_strip.lstrip('#').strip()

            if count_tokens(current_chunk + para) > config.MAX_TOKENS:
                if current_chunk.strip():
                    chunks.append({
                        'text': current_chunk.strip(),
                        'metadata': {
                            'source': file_path,
                            'source_name': Path(file_path).name,
                            'chunk_id': chunk_id,
                            'type': 'docx',
                            'section': current_section,
                            'indexed_at': datetime.now().isoformat()
                        }
                    })
                    chunk_id += 1
                current_chunk = para + "\n\n"
            else:
                current_chunk += para + "\n\n"
        
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'metadata': {
                    'source': file_path,
                    'source_name': Path(file_path).name,
                    'chunk_id': chunk_id,
                    'type': 'docx',
                    'section': current_section,
                    'indexed_at': datetime.now().isoformat()
                }
            })
        
        logger.info(f"✓ Processed docx {Path(file_path).name}, created {len(chunks)} chunks")
    except Exception as e:
        logger.error(f"Error processing docx {file_path}: {e}")
        raise DocumentProcessingError(f"Word file processing failed: {e}")
    return chunks

def process_file(file_info: Dict, company_id: str = None) -> List[Dict]:
    """Process file based on type, parsing frontmatter or sidecar metadata if present."""
    file_path = file_info['path']
    file_type = file_info['type']
    file_hash = file_info['hash']
    
    logger.info(f"📄 Processing {file_info['name']}...")
    
    from backend.src.excel_parser import process_excel_file, process_csv_file, load_tables_to_sqlite
    
    processors = {
        '.pdf': process_pdf,
        '.md': process_markdown,
        '.txt': process_text,
        '.docx': process_docx
    }
    
    chunks = []
    if file_type in ('.xlsx', '.xls'):
        file_chunks, sqlite_tables = process_excel_file(file_path)
        load_tables_to_sqlite(sqlite_tables, company_id)
        chunks = file_chunks
    elif file_type in ('.csv', '.tsv'):
        file_chunks, sqlite_tables = process_csv_file(file_path)
        load_tables_to_sqlite(sqlite_tables, company_id)
        chunks = file_chunks
    else:
        processor = processors.get(file_type)
        if processor:
            chunks = processor(file_path)
        else:
            logger.warning(f"Unsupported file type: {file_type}")
            return []

    # Load tag metadata from a sidecar JSON (binary files incl. xlsx/csv) or YAML
    # frontmatter (.md/.txt), then merge into every chunk — applies to all types.
    metadata_from_file = {}
    sidecar_path = file_path + ".metadata.json"
    if os.path.exists(sidecar_path):
        try:
            with open(sidecar_path, 'r', encoding='utf-8') as sf:
                metadata_from_file = json.load(sf)
            logger.info(f"Loaded sidecar metadata for {file_info['name']}")
        except Exception as e:
            logger.error(f"Error reading sidecar metadata {sidecar_path}: {e}")
    elif file_type in ['.md', '.txt']:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            match = re.match(r'^---\s*\n(.*?)\n---\s*\n', content, re.DOTALL)
            if match:
                import yaml
                frontmatter_text = match.group(1)
                metadata_from_file = yaml.safe_load(frontmatter_text) or {}
                logger.info(f"Loaded frontmatter metadata for {file_info['name']}")
        except Exception as e:
            logger.warning(f"Error reading frontmatter for {file_path}: {e}")

    for chunk in chunks:
        chunk['metadata']['file_hash'] = file_hash
        chunk['metadata']['relative_path'] = file_info.get('relative_path', '')
        for k, v in metadata_from_file.items():
            if k not in chunk['metadata']:
                chunk['metadata'][k] = v

    return chunks

# ---------------- RAG ENGINE CLASS ----------------
class RAGEngine:
    """PostgreSQL/pgvector-based RAG Engine."""

    def __init__(self, connection_uri: str = None, company_id: str = None, documents_dir: str = None):
        self.connection_uri = connection_uri or os.getenv("POSTGRES_URI") or os.getenv("DATABASE_URL")
        if not self.connection_uri:
            self.connection_uri = "postgresql://postgres:postgres@localhost:5432/argushr"
        # Multi-tenant: all rows are tagged with company_id and every query filters
        # on it, so tenants share the same tables but never see each other's data.
        self.company_id = company_id or "default"
        self.documents_dir = documents_dir or os.getenv("DOCUMENTS_DIR") or "documents"
        self.voyage_client = self._get_voyage_client()
        self.embed_dim = self._detect_embedding_dim()
        # psycopg2 connections are NOT safe to share across threads. This server
        # touches the DB from several threads at once (event loop, the folder
        # watcher's timer threads, and chat's threadpool). A single shared
        # connection corrupts under that concurrency ("invalid error value
        # specified"). Give each thread its own connection via thread-local storage.
        self._local = threading.local()
        self._index_lock = threading.Lock()
        self._init_db()

    @property
    def conn(self):
        """Return this thread's own PostgreSQL connection (created on first use)."""
        c = getattr(self._local, "conn", None)
        if c is not None and not getattr(c, "closed", 1):
            return c
        c = psycopg2.connect(self.connection_uri)
        c.autocommit = True
        self._local.conn = c
        return c

    def _close_thread_conn(self):
        c = getattr(self._local, "conn", None)
        if c is not None:
            try:
                c.close()
            except Exception:
                pass
            self._local.conn = None

    def _detect_embedding_dim(self) -> int:
        """
        Detect the embedding dimension from the model so the pgvector column size is
        always correct (voyage-4-large is 1024, not the old hard-coded 1536). Falls
        back to an env override, then 1024.
        """
        env_dim = os.getenv("EMBED_DIM")
        if env_dim and env_dim.isdigit():
            return int(env_dim)
        try:
            probe = self.voyage_client.embed(["dimension probe"], model=config.EMBED_MODEL)
            return len(probe.embeddings[0])
        except Exception as e:
            logger.warning(f"Could not probe embedding dim ({e}); defaulting to 1024.")
            return 1024

    def is_connected(self) -> bool:
        """Cheaply check whether this thread's connection is alive."""
        try:
            with self.conn.cursor() as cur:
                cur.execute("SELECT 1")
            return True
        except Exception:
            self._close_thread_conn()
            return False

    def reconnect(self):
        """Drop this thread's connection and re-run schema migrations."""
        self._close_thread_conn()
        self._init_db()

    def _get_voyage_client(self) -> voyageai.Client:
        """Get Voyage AI client."""
        key = os.getenv("VOYAGE_API_KEY2")
        if not key:
            raise ValueError("VOYAGE_API_KEY2 not set in environment")
        return voyageai.Client(api_key=key)
    
    def _init_db(self):
        """Run schema migrations on this thread's connection."""
        try:
            cursor = self.conn.cursor()
            
            # Enable pgvector extension
            try:
                cursor.execute("CREATE EXTENSION IF NOT EXISTS vector;")
            except Exception as ext_err:
                logger.warning(f"Could not create vector extension automatically: {ext_err}. Make sure it is installed and enabled.")
            
            # Create documents table. Multi-tenant: (company_id, filename) composite
            # PK so different companies can have same-named files without collision.
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS documents (
                    company_id VARCHAR(100) NOT NULL DEFAULT 'default',
                    filename VARCHAR(512) NOT NULL,
                    path TEXT NOT NULL,
                    hash VARCHAR(64) NOT NULL,
                    department VARCHAR(100),
                    uploaded_by VARCHAR(100),
                    category VARCHAR(100),
                    index_status VARCHAR(50) DEFAULT 'pending',
                    error_message TEXT,
                    vector_count INTEGER DEFAULT 0,
                    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    PRIMARY KEY (company_id, filename)
                );
            """)

            # Create chunks table with a matching composite FK + tenant column.
            cursor.execute(f"""
                CREATE TABLE IF NOT EXISTS chunks (
                    id VARCHAR(255) PRIMARY KEY,
                    company_id VARCHAR(100) NOT NULL DEFAULT 'default',
                    filename VARCHAR(512) NOT NULL,
                    chunk_index INTEGER NOT NULL,
                    content TEXT NOT NULL,
                    metadata JSONB NOT NULL,
                    embedding vector({self.embed_dim}) NOT NULL,
                    FOREIGN KEY (company_id, filename)
                        REFERENCES documents(company_id, filename) ON DELETE CASCADE
                );
            """)

            # Tenant lookup index (small, always useful).
            cursor.execute("CREATE INDEX IF NOT EXISTS chunks_company_file_idx ON chunks(company_id, filename);")

            # ANN index on embeddings is OPT-IN. HNSW is the single biggest disk
            # consumer in a pgvector store — big enough to fill a small volume and
            # crash the DB. At our scale (a few thousand chunks/company) an exact
            # cosine scan is fast and uses a fraction of the disk, so we default to
            # no ANN index. Set PGVECTOR_ANN_INDEX=hnsw (or ivfflat) to enable one
            # on a larger volume. Any leftover HNSW index from before is dropped
            # here so redeploying actually reclaims that space.
            ann = os.getenv("PGVECTOR_ANN_INDEX", "none").strip().lower()
            if ann == "hnsw":
                try:
                    cursor.execute("CREATE INDEX IF NOT EXISTS chunks_embedding_hnsw_idx ON chunks USING hnsw (embedding vector_cosine_ops);")
                    logger.info("ANN index: HNSW enabled.")
                except Exception as idx_err:
                    logger.warning(f"Could not create HNSW index ({idx_err}); using exact cosine search instead.")
            elif ann == "ivfflat":
                try:
                    cursor.execute("CREATE INDEX IF NOT EXISTS chunks_embedding_ivf_idx ON chunks USING ivfflat (embedding vector_cosine_ops) WITH (lists = 100);")
                    logger.info("ANN index: IVFFlat enabled.")
                except Exception as idx_err:
                    logger.warning(f"Could not create IVFFlat index ({idx_err}); using exact cosine search instead.")
            else:
                # Default: exact search, no ANN index. Drop any pre-existing HNSW
                # index so its disk is returned to the volume on next deploy.
                logger.info("ANN index: none (exact cosine search). Set PGVECTOR_ANN_INDEX=hnsw to enable.")
                for stale in ("chunks_embedding_hnsw_idx", "chunks_embedding_ivf_idx", "chunks_embedding_idx"):
                    try:
                        cursor.execute(f"DROP INDEX IF EXISTS {stale};")
                    except Exception as drop_err:
                        logger.debug(f"Could not drop {stale}: {drop_err}")
                
            logger.info(f"✅ PostgreSQL database initialized with connection.")
            
            cursor.execute("SELECT COUNT(*) FROM chunks")
            count = cursor.fetchone()[0]
            logger.info(f"   Collection has {count} chunks/vectors")
            
        except Exception as e:
            # Keep quiet here — the engine factory logs the down-state once and
            # throttles retries, so we don't spam a full traceback on every poll.
            logger.debug(f"PostgreSQL init failed: {e}")
            raise RAGError(f"Database initialization failed: {e}")
    
    def _embed_with_retry(self, texts: List[str], max_retries: int = 3) -> List[List[float]]:
        """Embed texts with exponential backoff retry."""
        all_embeddings = []
        
        # Create batches
        batches = []
        current_batch = []
        current_tokens = 0
        
        for text in texts:
            text_tokens = count_tokens(text)
            
            if (current_tokens + text_tokens > config.MAX_BATCH_TOKENS and current_batch) or \
               len(current_batch) >= config.MAX_BATCH_SIZE:
                batches.append(current_batch)
                current_batch = [text]
                current_tokens = text_tokens
            else:
                current_batch.append(text)
                current_tokens += text_tokens
        
        if current_batch:
            batches.append(current_batch)
        
        # Process batches
        for batch_num, batch in enumerate(batches, 1):
            retry_count = 0
            while retry_count < max_retries:
                try:
                    res = self.voyage_client.embed(batch, model=config.EMBED_MODEL)
                    all_embeddings.extend(res.embeddings)
                    break
                except Exception as e:
                    retry_count += 1
                    wait_time = min(60, 2 ** retry_count)
                    logger.warning(f"Embedding error (attempt {retry_count}): {e}")
                    if retry_count < max_retries:
                        logger.info(f"Retrying in {wait_time}s...")
                        time.sleep(wait_time)
                    else:
                        raise EmbeddingError(f"Failed after {max_retries} attempts: {e}")
            
            if batch_num < len(batches):
                time.sleep(config.WAIT_TIME)
        
        return all_embeddings
    
    def build_index(self, force_rebuild: bool = False) -> bool:
        """Serialize concurrent builds (the watcher can fire overlapping runs) and
        release this (often short-lived watcher) thread's connection afterwards."""
        with self._index_lock:
            try:
                return self._build_index_locked(force_rebuild)
            finally:
                self._close_thread_conn()

    def _build_index_locked(self, force_rebuild: bool = False) -> bool:
        """
        Build/update the index incrementally, STREAMING per file and per slice so
        peak process memory stays flat regardless of how many/large the files are.
        Previously every chunk + every 1024-dim embedding + every stringified
        vector for all new files was held in RAM at once, which OOM-killed the
        small Railway container on multi-file / large-spreadsheet ingests.

        A file that fails (parse error, oversized, embedding failure) is marked
        'failed' and skipped on future runs unless its bytes change — so one bad
        file can't OOM-loop the service on every restart.
        """
        logger.info("=" * 60)
        logger.info("📂 Discovering documents...")

        files = discover_files(self.documents_dir)
        if not files:
            logger.warning("No documents found. Please add files to documents/ folder")
            return False

        logger.info(f"Found {len(files)} file(s) in documents folder")

        cursor = self.conn.cursor()

        if force_rebuild:
            logger.info("🗑️ Clearing existing index...")
            try:
                cursor.execute("DELETE FROM documents WHERE company_id = %s", (self.company_id,))
                logger.info("✅ Database truncated.")
            except Exception as e:
                logger.error(f"Failed to clear index: {e}")

        # Successfully-indexed files (have chunks) → change detection.
        indexed_files = {}
        # Previously-failed files → skip if unchanged (avoids OOM retry loop).
        failed_files = {}
        if not force_rebuild:
            try:
                logger.info("🔍 Retrieving existing document metadata for change detection...")
                cursor.execute(
                    "SELECT c.filename, d.hash FROM chunks c "
                    "JOIN documents d ON c.filename = d.filename AND c.company_id = d.company_id "
                    "WHERE d.company_id = %s", (self.company_id,))
                for filename, file_hash in cursor.fetchall():
                    indexed_files.setdefault(filename, file_hash)
                logger.info(f"Loaded metadata for {len(indexed_files)} indexed file(s).")

                cursor.execute(
                    "SELECT filename, hash FROM documents "
                    "WHERE company_id = %s AND index_status = 'failed'", (self.company_id,))
                for filename, file_hash in cursor.fetchall():
                    failed_files[filename] = file_hash
                if failed_files:
                    logger.info(f"{len(failed_files)} previously-failed file(s) on record.")
            except Exception as e:
                logger.error(f"Error fetching metadata: {e}. Falling back to full rebuild.")
                indexed_files, failed_files = {}, {}

        current_files_by_path = {f['path']: f for f in files}
        files_to_delete = []
        files_to_process = []

        # Deleted files (were indexed, now gone from disk).
        # CRITICAL GUARD: Only delete from Neon if physical files still exist on disk
        # for this tenant. If current_files_by_path is empty (e.g. after a container
        # redeploy wipes the ephemeral disk), we MUST NOT delete anything — those
        # vectors in Neon are the only surviving copy of the data.
        if current_files_by_path:  # Only purge stale entries when disk has at least 1 file
            for indexed_path in indexed_files:
                if indexed_path not in current_files_by_path:
                    files_to_delete.append(indexed_path)
        else:
            logger.info("⚠️  No physical files found on disk — skipping stale-file purge to protect Neon data (ephemeral disk may have been wiped on redeploy).")

        # New or modified files.
        for f in files:
            path = f['path']
            current_hash = f['hash']
            if path in indexed_files:
                if current_hash != indexed_files[path]:
                    files_to_delete.append(path)
                    files_to_process.append(f)
                    logger.info(f"🔄 Modified file detected: {f['name']}")
                else:
                    logger.info(f"✅ Unchanged file skipped: {f['name']}")
            elif path in failed_files and failed_files[path] == current_hash:
                logger.warning(f"⏭️ Skipping previously-failed file (unchanged): {f['name']}")
            else:
                files_to_process.append(f)
                logger.info(f"🆕 New file detected: {f['name']}")

        # Deletions (connection is autocommit).
        if files_to_delete:
            logger.info("=" * 60)
            logger.info(f"🗑️ Deleting chunks for {len(files_to_delete)} removed/modified file(s)...")
            for path in files_to_delete:
                try:
                    cursor.execute("DELETE FROM documents WHERE company_id = %s AND filename = %s", (self.company_id, path))
                    logger.info(f"Deleted {Path(path).name} from PostgreSQL (cascade removed chunks)")
                except Exception as e:
                    logger.error(f"Failed to delete records for {path}: {e}")

        if not files_to_process:
            logger.info("=" * 60)
            logger.info("✅ No new or modified files to index. Database is up to date!")
            return True

        logger.info("=" * 60)
        logger.info(f"🔨 Processing {len(files_to_process)} new/modified file(s) (streaming)...")

        any_ok = False
        for file_info in files_to_process:
            try:
                if self._index_one_file(cursor, file_info):
                    any_ok = True
            except Exception as e:
                logger.error(f"❌ Failed to index {file_info['name']}: {e}", exc_info=True)
                try:
                    cursor.execute("ROLLBACK;")
                except Exception:
                    pass
                self._mark_file_failed(file_info['path'], file_info['hash'], str(e))

        logger.info("=" * 60)
        logger.info(f"✅ Index update complete. Total chunks: {self.get_chunk_count()}")
        return any_ok

    def _index_one_file(self, cursor, file_info: Dict) -> bool:
        """
        Chunk, embed, and store ONE file. Embeddings and inserts happen in slices
        of INDEX_SLICE_SIZE so only one slice's vectors are ever in memory. Runs in
        a single transaction so a mid-file failure leaves no partial chunks.
        Returns True if chunks were stored.
        """
        import uuid

        chunks = process_file(file_info, self.company_id)
        if not chunks:
            logger.warning(f"No chunks created from {file_info['name']}; nothing to store.")
            return False

        # Every chunk of one file shares the same source filename + hash.
        meta0 = chunks[0]['metadata']
        filename = meta0.get('source', file_info['path'])
        file_hash = meta0.get('file_hash', file_info['hash'])
        slice_size = getattr(config, "INDEX_SLICE_SIZE", 100)
        n = len(chunks)

        cursor.execute("BEGIN;")
        # Parent row first (FK target); 'pending' until every slice lands.
        cursor.execute("""
            INSERT INTO documents (company_id, filename, path, hash, department, uploaded_by, category, index_status, error_message, vector_count)
            VALUES (%s, %s, %s, %s, %s, %s, %s, 'pending', NULL, 0)
            ON CONFLICT (company_id, filename) DO UPDATE
            SET hash = EXCLUDED.hash, department = EXCLUDED.department, uploaded_by = EXCLUDED.uploaded_by,
                category = EXCLUDED.category, index_status = 'pending', error_message = NULL,
                vector_count = 0, updated_at = CURRENT_TIMESTAMP
        """, (self.company_id, filename, filename, file_hash,
              meta0.get('department', ''), meta0.get('uploaded_by', ''), meta0.get('category', '')))

        stored = 0
        for start in range(0, n, slice_size):
            part = chunks[start:start + slice_size]
            texts = [c['text'] for c in part]
            embeddings = self._embed_with_retry(texts)  # one bounded slice
            rows = []
            for c, emb in zip(part, embeddings):
                cid = f"chunk_{uuid.uuid4().hex[:16]}"
                embed_str = "[" + ",".join(map(str, emb)) + "]"
                rows.append((
                    cid, self.company_id, filename,
                    c['metadata'].get('chunk_index', c['metadata'].get('chunk_id', stored)),
                    c['text'], json.dumps(c['metadata'], default=str), embed_str,
                ))
                stored += 1
            execute_values(cursor, """
                INSERT INTO chunks (id, company_id, filename, chunk_index, content, metadata, embedding)
                VALUES %s
            """, rows)
            logger.info(f"  • {file_info['name']}: stored {stored}/{n} chunks")
            del texts, embeddings, rows  # free the slice before the next one

        cursor.execute("""
            UPDATE documents SET index_status = 'indexed', vector_count = %s, updated_at = CURRENT_TIMESTAMP
            WHERE company_id = %s AND filename = %s
        """, (stored, self.company_id, filename))
        cursor.execute("COMMIT;")
        logger.info(f"✅ Indexed {file_info['name']} ({stored} chunks).")
        del chunks
        return True

    def _mark_file_failed(self, filename: str, file_hash: str, error: str) -> None:
        """
        Record a file as 'failed' (autocommit) so future builds skip it unless its
        bytes change — a single bad/oversized file can't OOM-loop the service.
        """
        try:
            cur = self.conn.cursor()
            cur.execute("""
                INSERT INTO documents (company_id, filename, path, hash, department, uploaded_by, category, index_status, error_message, vector_count)
                VALUES (%s, %s, %s, %s, '', '', '', 'failed', %s, 0)
                ON CONFLICT (company_id, filename) DO UPDATE
                SET hash = EXCLUDED.hash, index_status = 'failed',
                    error_message = EXCLUDED.error_message, updated_at = CURRENT_TIMESTAMP
            """, (self.company_id, filename, filename, file_hash, (error or "")[:2000]))
            logger.warning(f"Marked '{Path(filename).name}' as failed; skipped until it changes.")
        except Exception as e:
            logger.error(f"Could not mark file failed: {e}")
    
    def query(self, query_text: str, top_k: int = None, use_llm: bool = True) -> Dict:
        """Query the RAG system."""
        top_k = top_k or config.RERANK_TOP_K
        
        if not self.conn:
            raise QueryError("Database connection is not initialized")
            
        try:
            cursor = self.conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM chunks WHERE company_id = %s", (self.company_id,))
            count = cursor.fetchone()[0]
        except Exception as e:
            raise QueryError(f"Error checking collection count: {e}")

        if count == 0:
            return {
                'answer': "No documents indexed. Please add documents to the documents/ folder.",
                'sources': []
            }

        # Embed query
        query_embedding = self._embed_with_retry([query_text])[0]

        # Search PostgreSQL using pgvector Cosine similarity — scoped to this tenant.
        try:
            query_embed_str = f"[{','.join(map(str, query_embedding))}]"
            cursor.execute("""
                SELECT content, metadata, (1 - (embedding <=> %s::vector)) AS score
                FROM chunks
                WHERE company_id = %s
                ORDER BY embedding <=> %s::vector
                LIMIT %s
            """, (query_embed_str, self.company_id, query_embed_str, config.TOP_K))
            rows = cursor.fetchall()
        except Exception as e:
            raise QueryError(f"Error searching PostgreSQL pgvector: {e}")
        
        if not rows:
            return {
                'answer': "No relevant documents found for your query.",
                'sources': []
            }
        
        # Prepare for reranking
        candidates = [row[0] for row in rows]
        metadatas = []
        for row in rows:
            meta = row[1]
            if isinstance(meta, str):
                try:
                    meta = json.loads(meta)
                except Exception:
                    pass
            metadatas.append(meta)
        
        # Rerank with Voyage
        try:
            rerank_results = self.voyage_client.rerank(
                query=query_text,
                documents=candidates,
                model=config.RERANK_MODEL,
                top_k=top_k
            )
        except Exception as e:
            logger.error(f"Reranking error: {e}. Falling back to top vector search results.")
            sources = []
            context_chunks = []
            for i in range(min(top_k, len(candidates))):
                sources.append({
                    'text': candidates[i],
                    'metadata': metadatas[i],
                    'score': float(rows[i][2])
                })
                context_chunks.append({
                    'text': candidates[i],
                    'metadata': metadatas[i]
                })
            
            answer = None
            if use_llm:
                answer = self._generate_answer(query_text, context_chunks)
            
            return {
                'answer': answer,
                'sources': sources
            }
        
        # Build sources from rerank results
        sources = []
        context_chunks = []
        for r in rerank_results.results:
            idx = r.index
            sources.append({
                'text': candidates[idx],
                'metadata': metadatas[idx],
                'score': r.relevance_score
            })
            context_chunks.append({
                'text': candidates[idx],
                'metadata': metadatas[idx]
            })
        
        # Generate answer
        answer = None
        if use_llm:
            answer = self._generate_answer(query_text, context_chunks)
        
        return {
            'answer': answer,
            'sources': sources
        }
    
    def _generate_answer(self, query: str, context_chunks: List[Dict]) -> str:
        """Generate answer using Groq."""
        try:
            from groq import Groq
        except ImportError:
            return "⚠️ groq not installed. Run: pip install groq"
        
        groq_key = os.getenv("GROQ_API_KEY")
        if not groq_key:
            return "⚠️ GROQ_API_KEY not set. Add it to your .env file."
        
        client = Groq(api_key=groq_key)
        
        # Build context
        current_tokens = 0
        context_parts = []
        
        for i, chunk in enumerate(context_chunks, 1):
            source_name = chunk['metadata'].get('source_name', 'Unknown')
            meta_info = ""
            
            chunk_type = chunk['metadata'].get('type')
            if chunk_type == 'excel':
                meta_info = f" (Sheet: {chunk['metadata'].get('sheet', 'N/A')})"
            elif chunk_type == 'pdf':
                meta_info = f" (Page: {chunk['metadata'].get('page', 'N/A')})"
            elif chunk_type == 'markdown':
                section = chunk['metadata'].get('section')
                if section:
                    meta_info = f" (Section: {section})"
            
            part_text = f"[Source {i}: {source_name}{meta_info}]\n{chunk['text']}"
            part_tokens = count_tokens(part_text)
            
            if current_tokens + part_tokens > config.MAX_CONTEXT_TOKENS:
                break
            
            context_parts.append(part_text)
            current_tokens += part_tokens
        
        context = "\n\n---\n\n".join(context_parts)
        
        prompt = f"""You are ArgusHR, an AI assistant specialized in Indian HR policies, employment regulations, and workplace compliance.
 
You help employees and HR professionals understand:
- Employment contracts and terms under Indian law
- Statutory benefits in India (EPF, Gratuity, ESI, Professional Tax)
- Regional state-level policies (Karnataka, Maharashtra, Delhi, Tamil Nadu)
- Company policies (BYOD, leave policy including SL, CL, PL, etc.)
- Employment types (Permanent, Contractor)
 
IMPORTANT RULES:
1. Answer ONLY based on the provided context
2. If the context doesn't contain the answer, say so clearly
3. Be specific and cite sources when possible
4. Use bullet points for clarity
5. Apply the correct state-level laws (e.g., Professional Tax slabs) if the user's state is specified
 
Context:
{context}
 
Question: {query}
 
Answer:"""
        
        # Try primary model with retries and cooldown delays
        for attempt in range(3):
            try:
                response = client.chat.completions.create(
                    model=config.LLM_MODEL,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.2,
                    max_tokens=2000,
                )
                return response.choices[0].message.content
            except Exception as e:
                err_str = str(e).lower()
                if "429" in err_str or "rate_limit" in err_str or "rate limit" in err_str:
                    if attempt < 2:
                        logger.warning(f"Rate limit hit on {config.LLM_MODEL} (attempt {attempt+1}/3). Cooldown sleeping 3s...")
                        time.sleep(3.0)
                        continue
                logger.error(f"LLM primary error: {e}")
                break
 
        # Primary failed, try fallback model
        logger.warning("Attempting backup fallback model llama-3.1-8b-instant...")
        for attempt in range(3):
            try:
                response = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.2,
                    max_tokens=2000,
                )
                return response.choices[0].message.content
            except Exception as e:
                err_str = str(e).lower()
                if "429" in err_str or "rate_limit" in err_str or "rate limit" in err_str:
                    if attempt < 2:
                        logger.warning(f"Rate limit hit on backup model llama-3.1-8b-instant (attempt {attempt+1}/3). Cooldown sleeping 3s...")
                        time.sleep(3.0)
                        continue
                logger.error(f"Fallback LLM error: {e}")
                return f"⚠️ Error generating response: {e}"
 
        return "⚠️ Error generating response: All attempts exhausted due to rate limits."
    
    def get_chunk_count(self) -> int:
        """Get number of chunks for this tenant."""
        if not self.conn:
            return 0
        try:
            cursor = self.conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM chunks WHERE company_id = %s", (self.company_id,))
            return cursor.fetchone()[0]
        except Exception:
            return 0

    def get_doc_count(self) -> int:
        """Get number of unique source documents for this tenant."""
        if not self.conn:
            return 0
        try:
            cursor = self.conn.cursor()
            cursor.execute("SELECT COUNT(*) FROM documents WHERE company_id = %s", (self.company_id,))
            return cursor.fetchone()[0]
        except Exception:
            return 0

    def get_counts(self) -> Tuple[int, int]:
        """Get (doc_count, chunk_count) in a single database roundtrip."""
        if not self.conn:
            return 0, 0
        try:
            cursor = self.conn.cursor()
            cursor.execute(
                "SELECT "
                "(SELECT COUNT(*) FROM documents WHERE company_id = %s), "
                "(SELECT COUNT(*) FROM chunks WHERE company_id = %s)",
                (self.company_id, self.company_id),
            )
            r = cursor.fetchone()
            return (r[0] if r else 0), (r[1] if r else 0)
        except Exception:
            return 0, 0

    # ---------------- BACKEND-AGNOSTIC HELPERS ----------------
    def list_indexed_documents(self) -> Dict[str, Dict]:
        """Return this tenant's indexed-document metadata keyed by base filename."""
        out: Dict[str, Dict] = {}
        if not self.conn:
            return out
        try:
            cursor = self.conn.cursor()
            cursor.execute(
                "SELECT filename, department, uploaded_by, category, index_status, "
                "error_message, vector_count, updated_at FROM documents WHERE company_id = %s",
                (self.company_id,),
            )
            for r in cursor.fetchall():
                base = os.path.basename(r[0])
                out[base] = {
                    "department": r[1],
                    "uploaded_by": r[2],
                    "category": r[3],
                    "index_status": r[4] or "indexed",
                    "error_message": r[5],
                    "vector_count": r[6] or 0,
                    "updated_at": r[7].isoformat() if isinstance(r[7], datetime) else str(r[7]) if r[7] else None,
                }
        except Exception as e:
            logger.error(f"list_indexed_documents failed: {e}")
        return out

    def get_chunks_for(self, source_name: str, sheet: str = None) -> List[Dict]:
        """
        Return all of this tenant's chunks for a file (optionally a sheet), ordered
        by chunk_index, so a whole sheet can be loaded into context for aggregation
        questions. Mirrors the ChromaDB engine's method.
        """
        if not self.conn:
            return []
        try:
            cursor = self.conn.cursor()
            cursor.execute(
                "SELECT content, metadata, chunk_index FROM chunks "
                "WHERE company_id = %s AND (filename = %s OR filename LIKE %s) "
                "ORDER BY chunk_index",
                (self.company_id, source_name, "%" + source_name),
            )
            items = []
            for content, meta, _ in cursor.fetchall():
                if isinstance(meta, str):
                    try:
                        meta = json.loads(meta)
                    except Exception:
                        meta = {}
                if sheet and (meta or {}).get("sheet") != sheet:
                    continue
                items.append({"text": content, "metadata": meta or {}})
            return items
        except Exception as e:
            logger.error(f"get_chunks_for failed for {source_name}/{sheet}: {e}")
            return []

    def delete_document(self, filename: str) -> int:
        """Delete this tenant's document + its chunks + its SQLite tables by filename."""
        try:
            from backend.src.excel_parser import delete_tables_from_sqlite
            delete_tables_from_sqlite(filename, self.company_id)
        except Exception as e:
            logger.error(f"Error dropping SQLite tables in delete_document: {e}")

        if not self.conn:
            return 0
        try:
            cursor = self.conn.cursor()
            cursor.execute(
                "DELETE FROM documents WHERE company_id = %s AND "
                "(filename = %s OR filename LIKE %s OR filename LIKE %s)",
                (self.company_id, filename, "%" + filename, "%" + filename.replace("\\", "/")),
            )
            removed = cursor.rowcount
            if not getattr(self.conn, "autocommit", False):
                self.conn.commit()
            return removed
        except Exception as e:
            logger.error(f"delete_document failed for {filename}: {e}")
            return 0

# ---------------- CLI INTERFACE ----------------
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="ArgusHR RAG Engine (PostgreSQL/pgvector)")
    parser.add_argument("--folder", "-f", type=str, default="documents",
                        help="Folder to index (e.g., 'documents/policies' or just 'policies')")
    parser.add_argument("--rebuild", "-r", action="store_true",
                        help="Force rebuild the index")
    parser.add_argument("--query", "-q", type=str,
                        help="Run a single query and exit")
    
    args = parser.parse_args()
    
    folder = args.folder
    if not folder.startswith("documents") and not os.path.isabs(folder):
        folder = f"documents/{folder}"
    Config.set_documents_dir(folder)
    
    print("=" * 60)
    print("🚀 ArgusHR RAG Engine (PostgreSQL/pgvector)")
    print("=" * 60)
    print(f"📂 Document folder: {config.DOCUMENTS_DIR}")
    
    try:
        engine = RAGEngine()
        engine.build_index(force_rebuild=args.rebuild)
    except Exception as e:
        logger.error(f"Initialization or indexing failed: {e}")
        exit(1)
    
    # Single query mode
    if args.query:
        print(f"\n🔍 Query: {args.query}\n")
        result = engine.query(args.query)
        print("=" * 60)
        print("🤖 Answer:")
        print("=" * 60)
        print(result['answer'])
        print("\n📚 Sources:")
        for i, src in enumerate(result['sources'], 1):
            print(f"  {i}. {src['metadata'].get('source_name', 'Unknown')} (score: {src['score']:.3f})")
        exit(0)
    
    # Interactive mode
    print("\n" + "=" * 60)
    print("RAG Query Interface")
    print("=" * 60)
    print("Commands: 'exit', 'rebuild', 'stats'")
    print()
    
    while True:
        try:
            user_input = input("❓ Your question: ").strip()
        except KeyboardInterrupt:
            print("\nGoodbye!")
            break
        
        if user_input.lower() in ["exit", "quit", ""]:
            print("Goodbye!")
            break
        
        if user_input.lower() == "rebuild":
            engine.build_index(force_rebuild=True)
            continue
        
        if user_input.lower() == "stats":
            print(f"\n📊 Index Statistics")
            print(f"   Chunks: {engine.get_chunk_count()}")
            print(f"   Documents: {engine.get_doc_count()}\n")
            continue
        
        print("\n🔍 Searching...\n")
        try:
            result = engine.query(user_input)
            print("=" * 60)
            print("🤖 Answer:")
            print("=" * 60)
            print(result['answer'])
            
            print("\n📚 Sources:")
            for i, src in enumerate(result['sources'], 1):
                print(f"  {i}. {src['metadata'].get('source_name', 'Unknown')} (score: {src['score']:.3f})")
            print()
        except Exception as e:
            print(f"❌ Query failed: {e}\n")
