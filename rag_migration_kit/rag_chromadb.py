"""
ArgusHR RAG Engine with ChromaDB
Enhanced RAG implementation with persistent vector storage.
"""
import os
import re
import sys
import json
import time
import hashlib
import logging
import threading

try:
    sys.stdout.reconfigure(encoding='utf-8')
except Exception:
    pass

from pathlib import Path
from typing import List, Dict, Tuple, Optional, Generator
from datetime import datetime

import numpy as np
import tiktoken
import voyageai
import chromadb
from chromadb.config import Settings
from dotenv import load_dotenv
import pandas as pd

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Load environment variables
load_dotenv(override=True)

# ---------------- CONFIGURATION ----------------
class Config:
    # Embedding settings
    EMBED_MODEL = "voyage-4-large"
    RERANK_MODEL = "rerank-2.5"
    
    # Chunking settings
    MAX_TOKENS = 800
    OVERLAP_TOKENS = 120
    
    # Retrieval settings
    TOP_K = 25
    RERANK_TOP_K = 3
    
    # API limits
    MAX_BATCH_SIZE = 100
    MAX_BATCH_TOKENS = 90000
    WAIT_TIME = 0.3

    # Indexing memory guard: embed + store this many chunks at a time so peak
    # process memory stays flat regardless of how many/large the files are.
    INDEX_SLICE_SIZE = 100

    # LLM settings
    MAX_CONTEXT_TOKENS = 12000  # Increased from 6000
    LLM_MODEL = "llama-3.3-70b-versatile"
    
    # Directories (can be overridden)
    DOCUMENTS_DIR = "documents"
    CHROMA_DIR = "chroma_db"
    
    @classmethod
    def set_documents_dir(cls, path: str):
        """Set documents directory to a specific path or subfolder."""
        cls.DOCUMENTS_DIR = path

config = Config()

# ---------------- TOKENIZER ----------------
enc = tiktoken.get_encoding("cl100k_base")

def count_tokens(text: str) -> int:
    """Count tokens in text."""
    return len(enc.encode(text))

# ---------------- CUSTOM EXCEPTIONS ----------------
class RAGError(Exception):
    """Base RAG exception."""
    pass

class DocumentProcessingError(RAGError):
    """Error processing a document."""
    pass

class EmbeddingError(RAGError):
    """Error generating embeddings."""
    pass

class QueryError(RAGError):
    """Error during query."""
    pass

# ---------------- FILE UTILITIES ----------------
def get_file_hash(file_path: str) -> str:
    """Generate MD5 hash for change detection."""
    hasher = hashlib.md5()
    try:
        with open(file_path, 'rb') as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hasher.update(chunk)
        return hasher.hexdigest()
    except IOError as e:
        logger.error(f"Could not hash file {file_path}: {e}")
        raise DocumentProcessingError(f"Cannot read file: {file_path}")

def discover_files(directory: str) -> List[Dict]:
    """Discover all supported files in directory."""
    supported_extensions = {'.xlsx', '.xls', '.pdf', '.md', '.txt', '.docx', '.csv', '.tsv'}
    files = []
    
    doc_path = Path(directory)
    if not doc_path.exists():
        logger.info(f"Creating {directory}/ folder...")
        doc_path.mkdir(parents=True, exist_ok=True)
        return []
    
    for file_path in doc_path.rglob('*'):
        if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
            try:
                relative_path = file_path.relative_to(doc_path)
            except ValueError:
                relative_path = file_path.name

            files.append({
                'path': str(file_path),
                'relative_path': str(relative_path),
                'name': file_path.name,
                'type': file_path.suffix.lower(),
                'hash': get_file_hash(str(file_path)),
                'size': file_path.stat().st_size,
                'modified': datetime.fromtimestamp(file_path.stat().st_mtime).isoformat()
            })
    
    return sorted(files, key=lambda x: x['name'])

# ---------------- DOCUMENT PROCESSORS ----------------
def _render_sheet_rows(df) -> List[str]:
    """
    Render a sheet as one text line per row, dropping empty cells. Read with
    header=None upstream so NO row (including the visual header row) is lost or
    mislabeled as pandas 'Unnamed:' columns. Fully-empty rows are skipped.
    """
    rows: List[str] = []
    for _, r in df.iterrows():
        cells = [str(v).strip() for v in r.tolist()]
        cells = [c for c in cells if c and c.lower() != "nan"]
        if cells:
            rows.append(" | ".join(cells))
    return rows


def _sheet_to_chunks(sheet_name: str, df, overlap_rows: int = 2) -> List[str]:
    """
    Chunk a sheet's rows into token-bounded, self-describing pieces. Each chunk is
    prefixed with the sheet name. Consecutive chunks OVERLAP by `overlap_rows` so a
    list/section that straddles a boundary still appears complete in at least one
    chunk. A single oversized row is hard-split by tokens so nothing is dropped.
    """
    rows = _render_sheet_rows(df)
    if not rows:
        return []

    prefix = f"## Sheet: {sheet_name}\n"
    budget = max(200, config.MAX_TOKENS - count_tokens(prefix))

    out: List[str] = []
    cur: List[str] = []
    cur_tokens = 0

    def flush(keep_overlap: bool):
        nonlocal cur, cur_tokens
        if cur:
            out.append(prefix + "\n".join(cur))
            cur = cur[-overlap_rows:] if (keep_overlap and overlap_rows) else []
            cur_tokens = sum(count_tokens(x) for x in cur)

    for row in rows:
        t = count_tokens(row)
        if t > budget:
            flush(False)
            toks = enc.encode(row)
            for s in range(0, len(toks), budget):
                out.append(prefix + enc.decode(toks[s:s + budget]))
            continue
        if cur_tokens + t > budget and cur:
            flush(True)
        cur.append(row)
        cur_tokens += t
    flush(False)
    return out


def process_excel(file_path: str) -> List[Dict]:
    """
    Process an Excel workbook: every sheet, every row. Reads with header=None so
    free-form sheets (merged cells, section layouts) aren't garbled, renders rows
    compactly, and chunks with overlap. Each chunk carries sheet + chunk_index so
    the whole sheet can be reassembled in order for aggregation questions.
    """
    chunks: List[Dict] = []
    try:
        with pd.ExcelFile(file_path) as excel_file:
            for sheet_name in excel_file.sheet_names:
                df = pd.read_excel(excel_file, sheet_name=sheet_name, header=None)
                df = df.dropna(how="all")  # drop fully-empty rows
                if df.empty:
                    continue

                parts = _sheet_to_chunks(sheet_name, df)
                for ci, part in enumerate(parts):
                    chunks.append({
                        "text": part,
                        "metadata": {
                            "source": file_path,
                            "source_name": Path(file_path).name,
                            "sheet": sheet_name,
                            "type": "excel",
                            "chunk_index": ci,
                            "indexed_at": datetime.now().isoformat(),
                        },
                    })
            logger.info(
                f"✓ Processed {len(excel_file.sheet_names)} sheet(s) from "
                f"{Path(file_path).name}, created {len(chunks)} chunks"
            )
    except Exception as e:
        logger.error(f"Error processing Excel {file_path}: {e}")
        raise DocumentProcessingError(f"Excel processing failed: {e}")

    return chunks

def process_pdf(file_path: str) -> List[Dict]:
    """Process PDF file with enhanced metadata."""
    chunks = []
    
    try:
        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)
            
            for page_num, page in enumerate(pdf.pages, 1):
                text = page.extract_text()
                
                if not text or len(text.strip()) < 50:
                    continue
                
                page_text = f"## Page {page_num} of {total_pages}\n\n{text}"
                
                if count_tokens(page_text) <= config.MAX_TOKENS:
                    chunks.append({
                        'text': page_text,
                        'metadata': {
                            'source': file_path,
                            'source_name': Path(file_path).name,
                            'page': page_num,
                            'total_pages': total_pages,
                            'type': 'pdf',
                            'indexed_at': datetime.now().isoformat()
                        }
                    })
                else:
                    paragraphs = text.split('\n\n')
                    current_chunk = f"## Page {page_num} of {total_pages}\n\n"
                    
                    for para in paragraphs:
                        if count_tokens(current_chunk + para) > config.MAX_TOKENS:
                            if current_chunk.strip():
                                chunks.append({
                                    'text': current_chunk.strip(),
                                    'metadata': {
                                        'source': file_path,
                                        'source_name': Path(file_path).name,
                                        'page': page_num,
                                        'total_pages': total_pages,
                                        'type': 'pdf',
                                        'indexed_at': datetime.now().isoformat()
                                    }
                                })
                            current_chunk = f"## Page {page_num} of {total_pages}\n\n{para}\n\n"
                        else:
                            current_chunk += para + "\n\n"
                    
                    if current_chunk.strip():
                        chunks.append({
                            'text': current_chunk.strip(),
                            'metadata': {
                                'source': file_path,
                                'source_name': Path(file_path).name,
                                'page': page_num,
                                'total_pages': total_pages,
                                'type': 'pdf',
                                'indexed_at': datetime.now().isoformat()
                            }
                        })
            
            logger.info(f"✓ Processed {total_pages} pages from {Path(file_path).name}, created {len(chunks)} chunks")
            
    except Exception as e:
        logger.error(f"Error processing PDF {file_path}: {e}")
        raise DocumentProcessingError(f"PDF processing failed: {e}")
    
    return chunks

def split_markdown_blocks(md: str) -> List[str]:
    """Split markdown into semantic blocks."""
    blocks = []
    buf = []
    in_code = False

    lines = md.splitlines()
    for line in lines:
        if line.strip().startswith("```"):
            in_code = not in_code
            buf.append(line)
            continue

        if in_code:
            buf.append(line)
            continue

        if re.match(r"^#{1,6}\s+", line):
            if buf:
                blocks.append("\n".join(buf).strip())
                buf = []
            buf.append(line)
            continue

        if "|" in line and re.match(r"^\s*\|.*\|\s*$", line):
            buf.append(line)
            continue

        if re.match(r"^\s*[-*+]\s+", line):
            buf.append(line)
            continue

        if line.strip() == "":
            if buf:
                blocks.append("\n".join(buf).strip())
                buf = []
            continue

        buf.append(line)

    if buf:
        blocks.append("\n".join(buf).strip())

    return [b for b in blocks if b.strip()]

def chunk_markdown_blocks(blocks: List[str], max_tokens: int, overlap_tokens: int) -> List[str]:
    """Chunk markdown blocks with overlap."""
    chunks = []
    current = []
    current_tokens = 0

    def flush_with_overlap():
        nonlocal current, current_tokens
        if not current:
            return
        chunk_text = "\n\n".join(current).strip()
        chunks.append(chunk_text)

        if overlap_tokens > 0:
            tokens = enc.encode(chunk_text)
            overlap = tokens[-overlap_tokens:] if len(tokens) > overlap_tokens else tokens
            overlap_text = enc.decode(overlap)
            current = [overlap_text]
            current_tokens = count_tokens(overlap_text)
        else:
            current = []
            current_tokens = 0

    for block in blocks:
        block_tokens = count_tokens(block)

        if block_tokens > max_tokens:
            if current:
                flush_with_overlap()

            tokens = enc.encode(block)
            start = 0
            while start < len(tokens):
                end = min(start + max_tokens, len(tokens))
                chunk = enc.decode(tokens[start:end])
                chunks.append(chunk)
                start = end - overlap_tokens if overlap_tokens > 0 and end - overlap_tokens > start else end
            current = []
            current_tokens = 0
            continue

        if current_tokens + block_tokens <= max_tokens:
            current.append(block)
            current_tokens += block_tokens
        else:
            flush_with_overlap()
            current.append(block)
            current_tokens += block_tokens

    if current:
        flush_with_overlap()

    return chunks

def process_markdown(file_path: str) -> List[Dict]:
    """Process Markdown file with header extraction."""
    chunks = []
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Extract document title
        title_match = re.search(r'^#\s+(.+)$', content, re.MULTILINE)
        doc_title = title_match.group(1) if title_match else Path(file_path).stem
        
        blocks = split_markdown_blocks(content)
        chunk_texts = chunk_markdown_blocks(blocks, config.MAX_TOKENS, config.OVERLAP_TOKENS)
        
        for i, text in enumerate(chunk_texts):
            # Extract section header if present
            header_match = re.search(r'^#{1,6}\s+(.+)$', text, re.MULTILINE)
            section = header_match.group(1) if header_match else None
            
            chunks.append({
                'text': text,
                'metadata': {
                    'source': file_path,
                    'source_name': Path(file_path).name,
                    'chunk_id': i,
                    'type': 'markdown',
                    'title': doc_title,
                    'section': section or "",
                    'indexed_at': datetime.now().isoformat()
                }
            })
        
        logger.info(f"✓ Processed {Path(file_path).name}, created {len(chunks)} chunks")
        
    except Exception as e:
        logger.error(f"Error processing Markdown {file_path}: {e}")
        raise DocumentProcessingError(f"Markdown processing failed: {e}")
    
    return chunks

def process_text(file_path: str) -> List[Dict]:
    """Process plain text file."""
    chunks = []
    
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        paragraphs = content.split('\n\n')
        current_chunk = ""
        chunk_id = 0
        
        for para in paragraphs:
            if count_tokens(current_chunk + para) > config.MAX_TOKENS:
                if current_chunk.strip():
                    chunks.append({
                        'text': current_chunk.strip(),
                        'metadata': {
                            'source': file_path,
                            'source_name': Path(file_path).name,
                            'chunk_id': chunk_id,
                            'type': 'text',
                            'indexed_at': datetime.now().isoformat()
                        }
                    })
                    chunk_id += 1
                current_chunk = para + "\n\n"
            else:
                current_chunk += para + "\n\n"
        
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'metadata': {
                    'source': file_path,
                    'source_name': Path(file_path).name,
                    'chunk_id': chunk_id,
                    'type': 'text',
                    'indexed_at': datetime.now().isoformat()
                }
            })
        
        logger.info(f"✓ Processed {Path(file_path).name}, created {len(chunks)} chunks")
        
    except Exception as e:
        logger.error(f"Error processing text {file_path}: {e}")
        raise DocumentProcessingError(f"Text processing failed: {e}")
    
    return chunks

def process_docx(file_path: str) -> List[Dict]:
    """Process Microsoft Word .docx file."""
    chunks = []
    try:
        import docx
        with open(file_path, "rb") as f:
            doc = docx.Document(f)
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text)
        
        # Also extract tables if present
        for table in doc.tables:
            if not table.rows:
                continue
            # Get headers from the first row of the table
            headers = [cell.text.strip().replace("\n", " ") for cell in table.rows[0].cells]
            headers_str = " | ".join([h for h in headers if h])
            
            # Process each row and append the headers context
            start_row = 1 if len(table.rows) > 1 else 0
            for row in table.rows[start_row:]:
                row_text = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                # Filter out empty rows
                if any(row_text):
                    # Format as: Table [Header1 | Header2]: Val1 | Val2
                    full_text.append(f"Table [{headers_str}]: " + " | ".join(row_text))
                    
        content = "\n\n".join(full_text)
        
        paragraphs = content.split('\n\n')
        current_chunk = ""
        chunk_id = 0
        current_section = ""
        
        for para in paragraphs:
            para_strip = para.strip()
            # Detect section heading pattern
            if re.match(r'^(?:SECTION\s+\d+|Section\s+\d+|\d+\.\d+|\d+\.\d+\.\d+)\b', para_strip) or \
               (para_strip.isupper() and len(para_strip) < 60) or \
               para_strip.startswith("###") or para_strip.startswith("##"):
                current_section = para_strip.lstrip('#').strip()

            if count_tokens(current_chunk + para) > config.MAX_TOKENS:
                if current_chunk.strip():
                    chunks.append({
                        'text': current_chunk.strip(),
                        'metadata': {
                            'source': file_path,
                            'source_name': Path(file_path).name,
                            'chunk_id': chunk_id,
                            'type': 'docx',
                            'section': current_section,
                            'indexed_at': datetime.now().isoformat()
                        }
                    })
                    chunk_id += 1
                current_chunk = para + "\n\n"
            else:
                current_chunk += para + "\n\n"
        
        if current_chunk.strip():
            chunks.append({
                'text': current_chunk.strip(),
                'metadata': {
                    'source': file_path,
                    'source_name': Path(file_path).name,
                    'chunk_id': chunk_id,
                    'type': 'docx',
                    'section': current_section,
                    'indexed_at': datetime.now().isoformat()
                }
            })
        
        logger.info(f"✓ Processed docx {Path(file_path).name}, created {len(chunks)} chunks")
    except Exception as e:
        logger.error(f"Error processing docx {file_path}: {e}")
        raise DocumentProcessingError(f"Word file processing failed: {e}")
    return chunks

def process_file(file_info: Dict, company_id: str = None) -> List[Dict]:
    """Process file based on type."""
    file_path = file_info['path']
    file_type = file_info['type']
    file_hash = file_info['hash']
    
    logger.info(f"📄 Processing {file_info['name']}...")
    
    from backend.src.excel_parser import process_excel_file, process_csv_file, load_tables_to_sqlite
    
    processors = {
        '.pdf': process_pdf,
        '.md': process_markdown,
        '.txt': process_text,
        '.docx': process_docx
    }
    
    chunks = []
    if file_type in ('.xlsx', '.xls'):
        file_chunks, sqlite_tables = process_excel_file(file_path)
        load_tables_to_sqlite(sqlite_tables, company_id)
        chunks = file_chunks
    elif file_type in ('.csv', '.tsv'):
        file_chunks, sqlite_tables = process_csv_file(file_path)
        load_tables_to_sqlite(sqlite_tables, company_id)
        chunks = file_chunks
    else:
        processor = processors.get(file_type)
        if processor:
            chunks = processor(file_path)

    if chunks:
        # Load tag metadata (department/category/uploaded_by) from a sidecar JSON
        # (binary files) or YAML frontmatter (.md/.txt), matching the pgvector engine.
        metadata_from_file = {}
        sidecar_path = file_path + ".metadata.json"
        if os.path.exists(sidecar_path):
            try:
                with open(sidecar_path, 'r', encoding='utf-8') as sf:
                    metadata_from_file = json.load(sf)
                logger.info(f"Loaded sidecar metadata for {file_info['name']}")
            except Exception as e:
                logger.error(f"Error reading sidecar metadata {sidecar_path}: {e}")
        elif file_type in ['.md', '.txt']:
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                match = re.match(r'^---\s*\n(.*?)\n---\s*\n', content, re.DOTALL)
                if match:
                    import yaml
                    metadata_from_file = yaml.safe_load(match.group(1)) or {}
                    logger.info(f"Loaded frontmatter metadata for {file_info['name']}")
            except Exception as e:
                logger.warning(f"Error reading frontmatter for {file_path}: {e}")

        for chunk in chunks:
            chunk['metadata']['file_hash'] = file_hash
            chunk['metadata']['relative_path'] = file_info.get('relative_path', '')
            for k, v in metadata_from_file.items():
                if k not in chunk['metadata'] and v is not None:
                    chunk['metadata'][k] = v
            # ChromaDB metadata values must be str/int/float/bool/None. Coerce
            # anything else (datetimes from YAML, list columns from Excel, etc.).
            chunk['metadata'] = _sanitize_metadata(chunk['metadata'])
        return chunks
    else:
        logger.warning(f"Unsupported file type: {file_type}")
        return []


def _sanitize_metadata(meta: Dict) -> Dict:
    """Coerce metadata into ChromaDB-safe primitives; drop None values."""
    clean = {}
    for k, v in meta.items():
        if v is None:
            continue
        if isinstance(v, (str, int, float, bool)):
            clean[k] = v
        else:
            clean[k] = str(v)
    return clean

# ---------------- RAG ENGINE CLASS ----------------
class RAGEngine:
    """ChromaDB-based RAG Engine."""
    
    def __init__(self, persist_directory: str = None, company_id: str = None, documents_dir: str = None):
        self.persist_directory = persist_directory or config.CHROMA_DIR
        # Multi-tenant isolation: each company gets its OWN Chroma collection, so
        # one tenant's vectors can never be returned to another.
        self.company_id = company_id or "default"
        self.collection_name = f"org_{self.company_id}_documents"
        self.documents_dir = documents_dir or config.DOCUMENTS_DIR
        self.voyage_client = self._get_voyage_client()
        self.chroma_client = None
        self.collection = None
        # Serializes concurrent index builds (upload trigger + folder watcher can
        # fire together) so the same new file isn't processed twice into duplicates.
        self._index_lock = threading.Lock()
        self._init_chroma()

    def _get_voyage_client(self) -> voyageai.Client:
        """Get Voyage AI client."""
        key = os.getenv("VOYAGE_API_KEY2")
        if not key:
            raise ValueError("VOYAGE_API_KEY2 not set in environment")
        return voyageai.Client(api_key=key)

    def _init_chroma(self):
        """Initialize ChromaDB for this company's isolated collection."""
        Path(self.persist_directory).mkdir(parents=True, exist_ok=True)

        self.chroma_client = chromadb.PersistentClient(
            path=self.persist_directory,
            settings=Settings(anonymized_telemetry=False)
        )

        # One collection per tenant.
        self.collection = self.chroma_client.get_or_create_collection(
            name=self.collection_name,
            metadata={"hnsw:space": "cosine"}
        )

        logger.info(f"✅ ChromaDB initialized at {self.persist_directory} (tenant='{self.company_id}')")
        logger.info(f"   Collection '{self.collection_name}' has {self.collection.count()} documents")
    
    def _embed_with_retry(self, texts: List[str], max_retries: int = 3) -> List[List[float]]:
        """Embed texts with exponential backoff retry."""
        all_embeddings = []
        
        # Create batches
        batches = []
        current_batch = []
        current_tokens = 0
        
        for text in texts:
            text_tokens = count_tokens(text)
            
            if (current_tokens + text_tokens > config.MAX_BATCH_TOKENS and current_batch) or \
               len(current_batch) >= config.MAX_BATCH_SIZE:
                batches.append(current_batch)
                current_batch = [text]
                current_tokens = text_tokens
            else:
                current_batch.append(text)
                current_tokens += text_tokens
        
        if current_batch:
            batches.append(current_batch)
        
        # Process batches
        for batch_num, batch in enumerate(batches, 1):
            retry_count = 0
            while retry_count < max_retries:
                try:
                    res = self.voyage_client.embed(batch, model=config.EMBED_MODEL)
                    all_embeddings.extend(res.embeddings)
                    break
                except Exception as e:
                    retry_count += 1
                    wait_time = min(60, 2 ** retry_count)  # Exponential backoff
                    logger.warning(f"Embedding error (attempt {retry_count}): {e}")
                    if retry_count < max_retries:
                        logger.info(f"Retrying in {wait_time}s...")
                        time.sleep(wait_time)
                    else:
                        raise EmbeddingError(f"Failed after {max_retries} attempts: {e}")
            
            if batch_num < len(batches):
                time.sleep(config.WAIT_TIME)
        
        return all_embeddings
    
    def build_index(self, force_rebuild: bool = False) -> bool:
        """Build or update the index incrementally using file hashes."""
        # Serialize concurrent builds so a double-trigger can't create duplicates.
        with self._index_lock:
            return self._build_index_locked(force_rebuild)

    def _build_index_locked(self, force_rebuild: bool = False) -> bool:
        logger.info("=" * 60)
        logger.info("📂 Discovering documents...")
        
        files = discover_files(self.documents_dir)
        
        if not files:
            logger.warning("No documents found. Please add files to documents/ folder")
            return False
        
        logger.info(f"Found {len(files)} file(s) in documents folder")
        
        # Clear existing data if rebuilding
        if force_rebuild and self.collection.count() > 0:
            logger.info("🗑️ Clearing existing index...")
            self.chroma_client.delete_collection(self.collection_name)
            self.collection = self.chroma_client.create_collection(
                name=self.collection_name,
                metadata={"hnsw:space": "cosine"}
            )
            logger.info("✅ Collection recreated.")
            
        # Get existing index metadata for incremental updates
        indexed_files = {} # Maps source path -> {'hash': file_hash, 'ids': [chunk_ids]}
        
        if not force_rebuild and self.collection.count() > 0:
            try:
                logger.info("🔍 Retrieving existing documents metadata for change detection...")
                existing_data = self.collection.get(include=["metadatas"])
                existing_ids = existing_data.get('ids', [])
                existing_metadatas = existing_data.get('metadatas', [])
                
                for cid, meta in zip(existing_ids, existing_metadatas):
                    if meta and 'source' in meta:
                        src = meta['source']
                        h = meta.get('file_hash')
                        if src not in indexed_files:
                            indexed_files[src] = {'hash': h, 'ids': []}
                        indexed_files[src]['ids'].append(cid)
                logger.info(f"Loaded metadata for {len(indexed_files)} indexed file(s).")
            except Exception as e:
                logger.error(f"Error fetching metadata from collection: {e}. Falling back to full rebuild.")
                indexed_files = {}
                
        # Identify changes
        current_files_by_path = {f['path']: f for f in files}
        
        files_to_delete = []  # List of source paths
        files_to_process = [] # List of file info dicts
        
        # 1. Detect deleted files (present in DB but not in documents folder)
        for indexed_path, info in indexed_files.items():
            if indexed_path not in current_files_by_path:
                files_to_delete.append(indexed_path)
                
        # 2. Detect new or modified files
        for f in files:
            path = f['path']
            current_hash = f['hash']
            
            if path not in indexed_files:
                # New file
                files_to_process.append(f)
                logger.info(f"🆕 New file detected: {f['name']}")
            else:
                indexed_hash = indexed_files[path]['hash']
                if current_hash != indexed_hash:
                    # Modified file
                    files_to_delete.append(path)
                    files_to_process.append(f)
                    logger.info(f"🔄 Modified file detected: {f['name']}")
                else:
                    logger.info(f"✅ Unchanged file skipped: {f['name']}")
                    
        # Perform deletions
        if files_to_delete:
            logger.info("=" * 60)
            logger.info(f"🗑️ Deleting chunks for {len(files_to_delete)} removed/modified file(s)...")
            for path in files_to_delete:
                if path in indexed_files:
                    ids_to_del = indexed_files[path]['ids']
                    try:
                        self.collection.delete(ids=ids_to_del)
                        logger.info(f"Deleted {len(ids_to_del)} chunks for {Path(path).name}")
                    except Exception as e:
                        logger.error(f"Failed to delete chunks for {path}: {e}")
                    
        # Process and embed new/modified files
        if not files_to_process:
            logger.info("=" * 60)
            logger.info("✅ No new or modified files to index. Database is up to date!")
            return True
            
        logger.info("=" * 60)
        logger.info(f"🔨 Processing {len(files_to_process)} new/modified file(s) (streaming)...")

        # Stream per file and per slice: only one slice's chunks + embeddings are
        # ever in memory, so many/large files can't spike RAM. A file that fails
        # is skipped (logged) rather than aborting the whole batch.
        import uuid
        slice_size = getattr(config, "INDEX_SLICE_SIZE", 100)
        any_ok = False

        for file_info in files_to_process:
            try:
                chunks = process_file(file_info, self.company_id)
            except DocumentProcessingError as e:
                logger.error(f"Skipping {file_info['name']}: {e}")
                continue
            except Exception as e:
                logger.error(f"❌ Failed to process {file_info['name']}: {e}", exc_info=True)
                continue

            if not chunks:
                logger.warning(f"No chunks created from {file_info['name']}; nothing to store.")
                continue

            n = len(chunks)
            stored = 0
            try:
                for start in range(0, n, slice_size):
                    part = chunks[start:start + slice_size]
                    texts = [c['text'] for c in part]
                    embeddings = self._embed_with_retry(texts)  # one bounded slice
                    ids = [f"chunk_{uuid.uuid4().hex[:16]}" for _ in part]
                    metadatas = [c['metadata'] for c in part]
                    self.collection.add(
                        ids=ids,
                        embeddings=embeddings,
                        documents=texts,
                        metadatas=metadatas,
                    )
                    stored += len(part)
                    logger.info(f"  • {file_info['name']}: stored {stored}/{n} chunks")
                    del texts, embeddings, ids, metadatas
                any_ok = True
                logger.info(f"✅ Indexed {file_info['name']} ({stored} chunks).")
            except Exception as e:
                logger.error(f"❌ Failed to embed/store {file_info['name']}: {e}", exc_info=True)
            finally:
                del chunks

        logger.info(f"✅ Index update complete. Total collection size: {self.collection.count()} chunks.")
        return any_ok
    
    def query(self, query_text: str, top_k: int = None, use_llm: bool = True) -> Dict:
        """Query the RAG system."""
        top_k = top_k or config.RERANK_TOP_K
        
        # Verify collection exists and is not stale
        try:
            count = self.collection.count()
        except Exception as e:
            if "does not exist" in str(e).lower() or "invalid collection" in str(e).lower():
                logger.info("Collection reference is stale/deleted. Re-fetching collection...")
                try:
                    self.collection = self.chroma_client.get_collection(self.collection_name)
                    count = self.collection.count()
                except Exception as inner_e:
                    raise QueryError(f"Failed to recover collection: {inner_e}")
            else:
                raise QueryError(f"Error checking collection count: {e}")
                
        if count == 0:
            return {
                'answer': "No documents indexed. Please add documents to the documents/ folder.",
                'sources': []
            }
        
        # Embed query
        query_embedding = self._embed_with_retry([query_text])[0]
        
        # Search ChromaDB
        try:
            results = self.collection.query(
                query_embeddings=[query_embedding],
                n_results=config.TOP_K,
                include=["documents", "metadatas", "distances"]
            )
        except Exception as e:
            # Fallback retry in case query itself fails due to stale reference
            if "does not exist" in str(e).lower() or "invalid collection" in str(e).lower():
                logger.info("Query failed. Re-fetching collection and retrying...")
                try:
                    self.collection = self.chroma_client.get_collection(self.collection_name)
                    results = self.collection.query(
                        query_embeddings=[query_embedding],
                        n_results=config.TOP_K,
                        include=["documents", "metadatas", "distances"]
                    )
                except Exception as inner_e:
                    raise QueryError(f"Failed to query after re-fetching collection: {inner_e}")
            else:
                raise QueryError(f"Error searching ChromaDB: {e}")
        
        if not results['documents'][0]:
            return {
                'answer': "No relevant documents found for your query.",
                'sources': []
            }
        
        # Prepare for reranking
        candidates = results['documents'][0]
        metadatas = results['metadatas'][0]
        
        # Rerank with Voyage
        rerank_results = self.voyage_client.rerank(
            query=query_text,
            documents=candidates,
            model=config.RERANK_MODEL,
            top_k=top_k
        )
        
        # Build sources
        sources = []
        context_chunks = []
        for r in rerank_results.results:
            idx = r.index
            sources.append({
                'text': candidates[idx],
                'metadata': metadatas[idx],
                'score': r.relevance_score
            })
            context_chunks.append({
                'text': candidates[idx],
                'metadata': metadatas[idx]
            })
        
        # Generate answer
        answer = None
        if use_llm:
            answer = self._generate_answer(query_text, context_chunks)
        
        return {
            'answer': answer,
            'sources': sources
        }
    
    def _generate_answer(self, query: str, context_chunks: List[Dict]) -> str:
        """Generate answer using Groq."""
        try:
            from groq import Groq
        except ImportError:
            return "⚠️ groq not installed. Run: pip install groq"
        
        groq_key = os.getenv("GROQ_API_KEY")
        if not groq_key:
            return "⚠️ GROQ_API_KEY not set. Add it to your .env file."
        
        client = Groq(api_key=groq_key)
        
        # Build context
        current_tokens = 0
        context_parts = []
        
        for i, chunk in enumerate(context_chunks, 1):
            source_name = chunk['metadata'].get('source_name', 'Unknown')
            meta_info = ""
            
            chunk_type = chunk['metadata'].get('type')
            if chunk_type == 'excel':
                meta_info = f" (Sheet: {chunk['metadata'].get('sheet', 'N/A')})"
            elif chunk_type == 'pdf':
                meta_info = f" (Page: {chunk['metadata'].get('page', 'N/A')})"
            elif chunk_type == 'markdown':
                section = chunk['metadata'].get('section')
                if section:
                    meta_info = f" (Section: {section})"
            
            part_text = f"[Source {i}: {source_name}{meta_info}]\n{chunk['text']}"
            part_tokens = count_tokens(part_text)
            
            if current_tokens + part_tokens > config.MAX_CONTEXT_TOKENS:
                break
            
            context_parts.append(part_text)
            current_tokens += part_tokens
        
        context = "\n\n---\n\n".join(context_parts)
        
        prompt = f"""You are ArgusHR, an AI assistant specialized in Indian HR policies, employment regulations, and workplace compliance.

You help employees and HR professionals understand:
- Employment contracts and terms under Indian law
- Statutory benefits in India (EPF, Gratuity, ESI, Professional Tax)
- Regional state-level policies (Karnataka, Maharashtra, Delhi, Tamil Nadu)
- Company policies (BYOD, leave policy including SL, CL, PL, etc.)
- Employment types (Permanent, Contractor)

IMPORTANT RULES:
1. Answer ONLY based on the provided context
2. If the context doesn't contain the answer, say so clearly
3. Be specific and cite sources when possible
4. Use bullet points for clarity
5. Apply the correct state-level laws (e.g., Professional Tax slabs) if the user's state is specified

Context:
{context}

Question: {query}

Answer:"""
        
        # Try primary model with retries and cooldown delays
        for attempt in range(3):
            try:
                response = client.chat.completions.create(
                    model=config.LLM_MODEL,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.2,
                    max_tokens=2000,
                )
                return response.choices[0].message.content
            except Exception as e:
                err_str = str(e).lower()
                if "429" in err_str or "rate_limit" in err_str or "rate limit" in err_str:
                    if attempt < 2:
                        logger.warning(f"Rate limit hit on {config.LLM_MODEL} (attempt {attempt+1}/3). Cooldown sleeping 3s...")
                        time.sleep(3.0)
                        continue
                logger.error(f"LLM primary error: {e}")
                break

        # Primary failed or was rate limited, try fallback model with retries and cooldown delays
        logger.warning("Attempting backup fallback model llama-3.1-8b-instant...")
        for attempt in range(3):
            try:
                response = client.chat.completions.create(
                    model="llama-3.1-8b-instant",
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.2,
                    max_tokens=2000,
                )
                return response.choices[0].message.content
            except Exception as e:
                err_str = str(e).lower()
                if "429" in err_str or "rate_limit" in err_str or "rate limit" in err_str:
                    if attempt < 2:
                        logger.warning(f"Rate limit hit on backup model llama-3.1-8b-instant (attempt {attempt+1}/3). Cooldown sleeping 3s...")
                        time.sleep(3.0)
                        continue
                logger.error(f"Fallback LLM error: {e}")
                return f"⚠️ Error generating response: {e}"

        return "⚠️ Error generating response: All attempts exhausted due to rate limits."
    
    def get_chunk_count(self) -> int:
        """Get number of chunks in collection."""
        return self.collection.count() if self.collection else 0
    
    def get_doc_count(self) -> int:
        """Get number of unique source documents."""
        if not self.collection or self.collection.count() == 0:
            return 0

        # Query all metadata to count unique sources
        results = self.collection.get(include=["metadatas"])
        sources = set(m.get('source', '') for m in results['metadatas'])
        return len(sources)

    # ---------------- BACKEND-AGNOSTIC HELPERS ----------------
    def is_connected(self) -> bool:
        """ChromaDB is embedded/local — healthy as long as the collection responds."""
        try:
            if self.collection is None:
                return False
            self.collection.count()
            return True
        except Exception:
            return False

    def reconnect(self):
        """Re-open the local ChromaDB store."""
        self.chroma_client = None
        self.collection = None
        self._init_chroma()

    def list_indexed_documents(self) -> Dict[str, Dict]:
        """
        Return indexed-document metadata keyed by base filename, derived from the
        collection's chunk metadata. Backend-agnostic shape shared with pgvector.
        """
        out: Dict[str, Dict] = {}
        try:
            if not self.collection or self.collection.count() == 0:
                return out
            data = self.collection.get(include=["metadatas"])
            for meta in data.get("metadatas", []):
                if not meta:
                    continue
                source = meta.get("source") or meta.get("source_name") or ""
                base = os.path.basename(source)
                if not base:
                    continue
                entry = out.setdefault(base, {
                    "department": meta.get("department"),
                    "uploaded_by": meta.get("uploaded_by"),
                    "category": meta.get("category"),
                    "index_status": "indexed",
                    "error_message": None,
                    "vector_count": 0,
                    "updated_at": meta.get("indexed_at"),
                })
                entry["vector_count"] += 1
        except Exception as e:
            logger.error(f"list_indexed_documents failed: {e}")
        return out

    def get_chunks_for(self, source_name: str, sheet: str = None) -> List[Dict]:
        """
        Return all chunks for a file (optionally a single sheet), ordered by
        sheet + chunk_index. Used to load a whole sheet into context so aggregation
        questions ("how many X", "list all Y") see every row, not just top-k.
        """
        try:
            if not self.collection or self.collection.count() == 0:
                return []
            if sheet:
                where = {"$and": [{"source_name": source_name}, {"sheet": sheet}]}
            else:
                where = {"source_name": source_name}
            data = self.collection.get(where=where, include=["documents", "metadatas"])
            items = list(zip(data.get("documents", []), data.get("metadatas", [])))
            items.sort(key=lambda x: (
                (x[1] or {}).get("sheet", ""),
                (x[1] or {}).get("chunk_index", 0),
            ))
            return [{"text": d, "metadata": m or {}} for d, m in items]
        except Exception as e:
            logger.error(f"get_chunks_for failed for {source_name}/{sheet}: {e}")
            return []

    def delete_document(self, filename: str) -> int:
        """Delete all chunks whose source matches (base) filename. Returns count removed."""
        try:
            from backend.src.excel_parser import delete_tables_from_sqlite
            delete_tables_from_sqlite(filename, self.company_id)
            
            if not self.collection or self.collection.count() == 0:
                return 0
            data = self.collection.get(include=["metadatas"])
            ids = data.get("ids", [])
            metas = data.get("metadatas", [])
            to_delete = [
                cid for cid, meta in zip(ids, metas)
                if meta and os.path.basename(meta.get("source", "")) == filename
            ]
            if to_delete:
                self.collection.delete(ids=to_delete)
            return len(to_delete)
        except Exception as e:
            logger.error(f"delete_document failed for {filename}: {e}")
            return 0

# ---------------- CLI INTERFACE ----------------
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(description="ArgusHR RAG Engine (ChromaDB)")
    parser.add_argument("--folder", "-f", type=str, default="documents",
                        help="Folder to index (e.g., 'documents/policies' or just 'policies')")
    parser.add_argument("--rebuild", "-r", action="store_true",
                        help="Force rebuild the index")
    parser.add_argument("--query", "-q", type=str,
                        help="Run a single query and exit")
    
    args = parser.parse_args()
    
    # Set documents directory
    folder = args.folder
    if not folder.startswith("documents") and not os.path.isabs(folder):
        folder = f"documents/{folder}"
    Config.set_documents_dir(folder)
    
    print("=" * 60)
    print("🚀 ArgusHR RAG Engine (ChromaDB)")
    print("=" * 60)
    print(f"📂 Document folder: {config.DOCUMENTS_DIR}")
    
    engine = RAGEngine()
    engine.build_index(force_rebuild=args.rebuild)
    
    # Single query mode
    if args.query:
        print(f"\n🔍 Query: {args.query}\n")
        result = engine.query(args.query)
        print("=" * 60)
        print("🤖 Answer:")
        print("=" * 60)
        print(result['answer'])
        print("\n📚 Sources:")
        for i, src in enumerate(result['sources'], 1):
            print(f"  {i}. {src['metadata'].get('source_name', 'Unknown')} (score: {src['score']:.3f})")
        exit(0)
    
    # Interactive mode
    print("\n" + "=" * 60)
    print("RAG Query Interface")
    print("=" * 60)
    print("Commands: 'exit', 'rebuild', 'stats'")
    print()
    
    while True:
        user_input = input("❓ Your question: ").strip()
        
        if user_input.lower() in ["exit", "quit", ""]:
            print("Goodbye!")
            break
        
        if user_input.lower() == "rebuild":
            engine.build_index(force_rebuild=True)
            continue
        
        if user_input.lower() == "stats":
            print(f"\n📊 Index Statistics")
            print(f"   Chunks: {engine.get_chunk_count()}")
            print(f"   Documents: {engine.get_doc_count()}\n")
            continue
        
        print("\n🔍 Searching...\n")
        result = engine.query(user_input)
        
        print("=" * 60)
        print("🤖 Answer:")
        print("=" * 60)
        print(result['answer'])
        
        print("\n📚 Sources:")
        for i, src in enumerate(result['sources'], 1):
            print(f"  {i}. {src['metadata'].get('source_name', 'Unknown')} (score: {src['score']:.3f})")
        print()

